"""
Универсальный парсер файлов для Knowledge Base Pipeline.

Поддержка: PDF, FB2, DOCX, ODT, TXT
FB2 — приоритетный формат (парсинг по XML-тегам <p>, <title>).
"""
import re
import logging
from pathlib import Path
from typing import Optional

logger = logging.getLogger("kb.parsers")


def extract_text(file_path: Path) -> Optional[str]:
    """
    Универсальная точка входа: определяет формат и вызывает нужный парсер.

    Returns:
        Чистый текст или None при ошибке
    """
    suffix = file_path.suffix.lower()

    parsers = {
        ".pdf": _parse_pdf,
        ".fb2": _parse_fb2,
        ".txt": _parse_txt,
        ".docx": _parse_docx,
        ".odt": _parse_odt,
    }

    parser = parsers.get(suffix)
    if not parser:
        logger.warning(f"⚠️ Неподдерживаемый формат: {suffix} ({file_path.name})")
        return None

    try:
        text = parser(file_path)
        if text:
            text = _clean_text(text)
        return text if text and len(text) > 50 else None
    except Exception as e:
        logger.error(f"❌ Ошибка парсинга {file_path.name}: {e}")
        return None


# ============================================================
# PDF — PyMuPDF (fitz)
# ============================================================

def _parse_pdf(file_path: Path) -> str:
    """Извлечь текст из PDF через PyMuPDF"""
    import fitz  # PyMuPDF

    doc = fitz.open(str(file_path))
    text = ""
    for page_num in range(len(doc)):
        page = doc[page_num]
        text += page.get_text()
    doc.close()

    return text


# ============================================================
# FB2 — XML парсинг (приоритетный формат)
# ============================================================

def _parse_fb2(file_path: Path) -> str:
    """
    Извлечь текст из FB2 по XML-тегам <p>, <title>, <section>.
    FB2 — приоритетный формат для сохранения семантической структуры абзацев.
    """
    import xml.etree.ElementTree as ET

    # Читаем файл с автодетекцией кодировки
    raw_bytes = file_path.read_bytes()
    content = _decode_bytes(raw_bytes)

    # FB2 namespace
    # Пробуем найти namespace из корневого тега
    ns_match = re.search(r'xmlns="([^"]+)"', content[:1000])
    ns = {"fb": ns_match.group(1)} if ns_match else {}

    try:
        root = ET.fromstring(content)
    except ET.ParseError:
        # Иногда FB2 содержит невалидный XML, пробуем без namespace
        content = re.sub(r'xmlns="[^"]+"', '', content)
        root = ET.fromstring(content)

    parts = []

    # Извлекаем текст из <body>
    for body in root.iter():
        tag = _strip_ns(body.tag)

        if tag == "title":
            # Заголовки разделов
            title_text = _get_all_text(body)
            if title_text.strip():
                parts.append(f"\n## {title_text.strip()}\n")

        elif tag == "p":
            # Абзацы
            p_text = _get_all_text(body)
            if p_text.strip():
                parts.append(p_text.strip())

        elif tag == "empty-line":
            parts.append("")

    return "\n".join(parts)


def _strip_ns(tag: str) -> str:
    """Убирает namespace из XML-тега: {http://...}p → p"""
    if "}" in tag:
        return tag.split("}", 1)[1]
    return tag


def _get_all_text(element) -> str:
    """Рекурсивно извлекает весь текст из XML-элемента, включая вложенные теги"""
    texts = []
    if element.text:
        texts.append(element.text)
    for child in element:
        texts.append(_get_all_text(child))
        if child.tail:
            texts.append(child.tail)
    return " ".join(texts)


# ============================================================
# TXT — Plain text с автодетекцией кодировки
# ============================================================

def _parse_txt(file_path: Path) -> str:
    """Прочитать TXT с автоматической детекцией кодировки"""
    raw_bytes = file_path.read_bytes()
    return _decode_bytes(raw_bytes)


# ============================================================
# DOCX — python-docx
# ============================================================

def _parse_docx(file_path: Path) -> str:
    """Извлечь текст из DOCX"""
    from docx import Document

    doc = Document(str(file_path))
    parts = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            # Сохраняем заголовки с маркером
            if para.style and para.style.name.startswith("Heading"):
                level = para.style.name.replace("Heading ", "").replace("Heading", "1")
                try:
                    hashes = "#" * int(level)
                except ValueError:
                    hashes = "##"
                parts.append(f"\n{hashes} {text}\n")
            else:
                parts.append(text)

    # Таблицы
    for table in doc.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip():
                parts.append(row_text)

    return "\n".join(parts)


# ============================================================
# ODT — odfpy
# ============================================================

def _parse_odt(file_path: Path) -> str:
    """
    Извлечь текст из ODT (OpenDocument Text).

    Использует zipfile + ElementTree напрямую (ODT — это ZIP с content.xml),
    чтобы обойти ошибку ExternalReferenceForbidden в lxml/odfpy.
    """
    import zipfile
    import xml.etree.ElementTree as ET

    # ODT — это ZIP-архив
    try:
        with zipfile.ZipFile(str(file_path), "r") as zf:
            # Основной контент в content.xml
            if "content.xml" not in zf.namelist():
                raise ValueError("content.xml не найден в ODT-архиве")
            xml_bytes = zf.read("content.xml")
    except zipfile.BadZipFile:
        raise ValueError(f"Файл не является валидным ODT/ZIP: {file_path.name}")

    # Парсинг XML без загрузки внешних DTD
    root = ET.fromstring(xml_bytes)

    # ODT namespace для text:p (абзацы) и text:h (заголовки)
    NS = {
        "text": "urn:oasis:names:tc:opendocument:xmlns:text:1.0",
    }

    parts = []

    for elem in root.iter():
        # Убираем namespace из тега для сравнения
        local_tag = elem.tag.split("}")[-1] if "}" in elem.tag else elem.tag

        if local_tag in ("p", "h"):
            # Собираем весь текст из элемента и дочерних
            text_content = "".join(elem.itertext()).strip()
            if text_content:
                if local_tag == "h":
                    parts.append(f"\n## {text_content}\n")
                else:
                    parts.append(text_content)

    return "\n".join(parts)


def _odt_get_text(element) -> str:
    """Рекурсивно извлечь текст из ODT-элемента (legacy, для совместимости)"""
    result = ""
    for node in element.childNodes:
        if hasattr(node, "data"):
            result += node.data
        else:
            result += _odt_get_text(node)
    return result


# ============================================================
# Утилиты
# ============================================================

def _decode_bytes(raw_bytes: bytes) -> str:
    """Декодировать байты с автодетекцией кодировки"""
    # Пробуем UTF-8 первым (наиболее частый)
    try:
        return raw_bytes.decode("utf-8")
    except UnicodeDecodeError:
        pass

    # Пробуем cp1251 (Windows-кириллица)
    try:
        return raw_bytes.decode("cp1251")
    except UnicodeDecodeError:
        pass

    # Фолбэк — chardet
    try:
        import chardet
        detected = chardet.detect(raw_bytes)
        encoding = detected.get("encoding", "utf-8")
        return raw_bytes.decode(encoding, errors="replace")
    except ImportError:
        return raw_bytes.decode("utf-8", errors="replace")


def _clean_text(text: str) -> str:
    """Общая чистка текста после парсинга"""
    # Убираем лишние переносы строк (3+ → 2)
    text = re.sub(r'\n{3,}', '\n\n', text)
    # Убираем лишние пробелы/табы
    text = re.sub(r'[ \t]{2,}', ' ', text)
    # Убираем одиночные номера страниц на отдельных строках
    text = re.sub(r'^\d{1,4}$', '', text, flags=re.MULTILINE)
    # Убираем BOM и нулевые символы
    text = text.replace('\ufeff', '').replace('\x00', '')
    return text.strip()
